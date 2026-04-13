import logging
import re
from pathlib import Path

from docx import Document

from app.core.config import settings
from app.core.database import get_client
from app.services.log import log_event
from app.services.notification import notify

logger = logging.getLogger(__name__)

def _safe_filename(text: str) -> str:
    """Strip characters that are unsafe in filenames and truncate to 80 chars."""
    return re.sub(r"[^\w\s-]", "", text).strip().replace(" ", "_")[:80]


def compile_book(book_id: str) -> str:
    """
    Assemble all approved chapters into .docx and .txt files.

    Output is written to:  <OUTPUT_DIR>/<book_title>/

    Updates books.output_path and sets books.final_status = 'completed'.
    Returns the output directory path as a string.
    Raises ValueError if the book or its chapters cannot be found.
    """
    db = get_client()

    bk_resp = db.table("books").select("title").eq("id", book_id).execute()
    if not bk_resp.data:
        raise ValueError(f"Book {book_id!r} not found")
    book_title: str = bk_resp.data[0]["title"]

    ch_resp = (
        db.table("chapters")
        .select("chapter_index, title, content")
        .eq("book_id", book_id)
        .eq("status", "approved")
        .order("chapter_index")
        .execute()
    )
    chapters = ch_resp.data or []

    if not chapters:
        raise ValueError(f"No approved chapters found for book {book_id!r}")

    output_dir = Path(settings.output_dir) / _safe_filename(book_title)
    output_dir.mkdir(parents=True, exist_ok=True)
    base_name = _safe_filename(book_title)

    # ── .docx ──
    docx_path = output_dir / f"{base_name}.docx"
    doc = Document()
    doc.add_heading(book_title, level=0)

    for chapter in chapters:
        ch_title = chapter.get("title") or f"Chapter {chapter['chapter_index']}"
        doc.add_heading(
            f"Chapter {chapter['chapter_index']}: {ch_title}", level=1
        )
        doc.add_paragraph(chapter.get("content") or "")

    doc.save(str(docx_path))
    logger.info("Wrote %s", docx_path)

    # ── .txt ──
    txt_path = output_dir / f"{base_name}.txt"
    lines: list[str] = [book_title, "=" * len(book_title), ""]

    for chapter in chapters:
        ch_title = chapter.get("title") or f"Chapter {chapter['chapter_index']}"
        heading = f"Chapter {chapter['chapter_index']}: {ch_title}"
        lines += [heading, "-" * len(heading), "", chapter.get("content") or "", ""]

    txt_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("Wrote %s", txt_path)

    # ── persist & notify ──
    output_path = str(output_dir)

    db.table("books").update({
        "output_path": output_path,
        "final_status": "completed",
    }).eq("id", book_id).execute()

    log_event(
        "book_compiled",
        f"\"{book_title}\" compiled to {output_path}",
        book_id=book_id,
    )
    notify(
        "compilation_done",
        book_title=book_title,
        details=f"Your book has been compiled and saved to:\n{output_path}",
    )

    return output_path
